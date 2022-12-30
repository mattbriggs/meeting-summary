import docx
import sqlite3
import html

import sentiment as SE
import createdb as DB
import summarize as SU


def load_table_line(row, dbpath):
    '''Load the document table with a processDocument object.'''
    try:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        cur.execute('INSERT INTO line (ID, TStamp, Speaker, \
            Verbatim, SENT_POS, SENT_NEU, SENT_NEG, Sentiment) VALUES \
            ( ?, ?, ?, ?, ?, ?, ?, ?)', \
            ( row[0], row[1],  row[2],  row[3],  row[4],  row[5],  row[6],  row[7]) )
        conn.commit()
        cur.close()
    except Exception as e:
        print("An error occurred in loadDocument/document {}".format(e))


def table_transcript(filein, dbpath, size=0):
    '''With a filepath to a transcript, and an optional line index, return a table.'''

    doc = docx.Document(filein)
    if size == 0:
        size = len(doc.paragraphs)

    for indx, i in enumerate(doc.paragraphs):
        if indx > 3 and indx < size:
            if i.text.find("Speaker ") > 0:
                parsed = i.text.split(" ")
            else:
                print("Line: {}".format(size-indx))
                esctext = html.escape(i.text)
                sent = SE.get_sentiment(i.text)
                row = [indx, parsed[0], parsed[1] + " " + parsed[2], esctext, sent['neg'], sent['neu'], sent['pos'], sent['compound']]
                load_table_line(row, dbpath)

def get_verbatims(dbpath):
    '''With a loaded DB, collect varbitms.'''
    conn = sqlite3.connect(dbpath)
    cur = conn.cursor()
    speaker_list = list(cur.execute('SELECT * From speaker_list;'))
    verbatims = list(cur.execute('SELECT * From verbatims;'))
    cur.close()

    # get speakers
    verbatims_speaks = {}
    for i in speaker_list:
        verbatims_speaks[i[0]] = ""
    
    # prep long summaries
    summaries = {}
    for i in speaker_list:
        summaries[i[0]] = ""
    
    # prep short summaries
    shorts = {}
    for i in speaker_list:
        shorts[i[0]] = ""
    
    # get verbatims
    for i in speaker_list:
        for j in verbatims:
            if j[0] == i[0]:
                extract = str(j[1])
                verbatims_speaks[i[0]] += extract
    
    # get long summary of verbatims
    for i in speaker_list:
        summary = SU.summarize_text(html.unescape(verbatims_speaks[i[0]]))
        summaries[i[0]] = summary
        short = SU.summarize_text(html.unescape(verbatims_speaks[i[0]]), 400)
        shorts[i[0]] = short
    
    # get short summary and store all summary fields
    for i in speaker_list:
        conn = sqlite3.connect(dbpath)
        cur = conn.cursor()
        cur.execute('INSERT INTO summary (Speaker, ShortSummary, LongSummary, Allsaid) \
            VALUES ( ?, ?, ?, ?)', \
            ( i[0], shorts[i[0]], summaries[i[0]], verbatims_speaks[i[0]]))
        conn.commit()
        cur.close()

dbpath = "wordsample/summary.db"

db  = DB.DBModel()
db.create(dbpath, 'meetingsummary.sql')
table_transcript('wordsample/2022.11.28.MartinandPryianka.docx', dbpath)
get_verbatims(dbpath)