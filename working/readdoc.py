import docx
import csv
import html


def write_csv(outbody, path):
    '''Write CSV file to the path.'''
    csvout = open(path, 'w', newline="")
    csvwrite = csv.writer(csvout)
    for r in outbody:
        try:
            csvwrite.writerow(r)
        except Exception as e:
            print("An error: {}".format(e))
    csvout.close()

def table_transcript(filein, size=0):
    '''With a filepath to a transcript, and an optional line index, return a table.'''

    doc = docx.Document(filein)
    if size == 0:
        size = len(doc.paragraphs)

    table = [["time", "speaker", "line"]]

    for indx, i in enumerate(doc.paragraphs):
        if indx > 3 and indx < size:
            if i.text.find("Speaker ") > 0:
                parsed = i.text.split(" ")
            else:
                esctext = html.escape(i.text)
                table.append([parsed[0], parsed[1] + " " + parsed[2], esctext])
        
    return table


get_table = table_transcript('wordsample/2022.11.28.MartinandPryianka.docx')
write_csv(get_table, 'wordsample/tableoftrascript.csv')

